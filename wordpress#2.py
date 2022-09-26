
from string import hexdigits
import sys
from turtle import heading
sys.path.append("E:\code\python\Lib\site-packages")
from wordpress_xmlrpc import Client, WordPressPost
from wordpress_xmlrpc.methods import posts
import docx
import glob
import tkinter as tk
from tkinter import *
from tkinter import Tk,Scrollbar,Frame,RIGHT,Y,LEFT,messagebox,scrolledtext
from tkinter.ttk import Treeview
from pathlib import Path
from tkinter.messagebox import *

#输出文本内容
def print_text(file):
    text1 = "段落数:"+str(len(file.paragraphs))
    for para in file.paragraphs:
        text1 = (text1 + '\n'+ para.text)
    return text1

def text2(file):
    n=0
    text1 = "段落数:"+str(len(file.paragraphs))
    for para in file.paragraphs:
        text1 = (text1 + '\n'+ para.text)
        n=n+1
        if n>=3:
            break
#发布
def post_private(title, content,category,post_tag):
    post = WordPressPost()
    post.title = title
    post.content = content
    post.post_status = 'provate' #私密
    post.terms_names = {
        'category': [category],
        'post_tag': [post_tag],
    }
    post.id = wp.call(posts.NewPost(post))
    return post.id

def post_draft(title, content,category,post_tag):
    post = WordPressPost()
    post.title = title
    post.content = content
    post.post_status = 'draft'  #草稿
    post.terms_names = {
        'category': [category],
        'post_tag': [post_tag],
    }
    post.id = wp.call(posts.NewPost(post))
    return post.id

def post_publish(title, content,category,post_tag):
    post = WordPressPost()
    post.title = title
    post.content = content
    post.post_status = 'publish'  #公开
    post.terms_names = {
        'category': [category],
        'post_tag': [post_tag],
    }
    post.id = wp.call(posts.NewPost(post))
    return post.id

#模式显示
def print_selection():
    l.config(text= var.get())

#发送指令
def post_new_article():
    cat = category.get()
    tag = post_tag.get()
    file = art[art_no]
    file2=docx.Document(file)
    p=Path(file)
    global wp
    address1 = address.get()
    Id1 = Id.get()
    password1 = password.get()
    print(address1,Id1 ,password1)
    wp = Client(address1,password1,Id1)
    if var.get() == '私密':
        post_id = post_private(p.stem,print_text(file2),cat,tag) 
        mx1 = showinfo(title='提示', message=p.stem+'发送成功'+'ID:'+post_id)
    elif var.get() == '草稿':
        post_id = post_draft(p.stem,print_text(file2),cat,tag) 
        mx1 = showinfo(title='提示', message=p.stem+'发送成功'+'ID:'+post_id)
    elif var.get() == '公开':
        post_id = post_publish(p.stem,print_text(file2),cat,tag) 
        mx1 = showinfo(title='提示', message=p.stem+'发送成功'+'ID:'+post_id)
    else :
        print('错误')

if __name__ == '__main__':
    #导入文件
    art = glob.glob(r'./article/*')

    #创造面板
    root = tk.Tk()
    root.geometry('1700x800+100+100')
    root.title('WordPress管理系统')

    #创建容器
    frame = Frame(root)
    frame.place(x=0,y=50,width=370,height=280)

    frame2 = Frame(root)
    frame2.place(x=200,y=20,width=1024,height=20)

    frame3 = Frame(root)
    frame3.place(x=400,y=60,width=1024,height=720)

    frame4 = Frame(root)
    frame4.place(x=0,y=400,width=370,height=440)

    number_var = tk.StringVar() 
    number_var.set('请选择文章')
    number_var2 = scrolledtext.ScrolledText(frame3,width=100,height=30,font=("隶书",15)) 
    number_var2.place(x=0,y=50)
    i = '这里会显示你的文章内容'
    number_var2.insert(END,i)

    scrollBar = Scrollbar(frame)
    scrollBar.pack(side=RIGHT,fill=Y)

    tree = Treeview(frame,
                columns=('1','2'),
                show='headings',
                yscrollcommand=scrollBar.set)

    tk.Label(frame3,textvariable=number_var,bd=1,relief=tk.SUNKEN,anchor=tk.CENTER).pack(fill=tk.X)   

    #建立表格
    tree.column('1',width=50,anchor='center')
    tree.column('2',width=300,anchor='center')
    tree.heading('1',text='序号')
    tree.heading('2',text='文件名')

    tree.pack(side=LEFT,fill=Y)
    scrollBar.config(command=tree)

    for name in art:
        name2 = Path(name)
        title = name2.stem
        x = (art.index(name)+1, title)
        tree.insert('','end',values=x)
    
    labe1 = Label(frame2,text='地址')
    labe1.pack(side=LEFT) 
    address = tk.Entry(frame2,width=100)
    address.pack(side=LEFT) 

    labe2 = Label(frame2,text='用户名')
    labe2.pack(side=LEFT) 
    Id = tk.Entry(frame2,width=20,show='*')
    Id.pack(side=RIGHT)

    labe3 = Label(frame2,text='密码')
    labe3.pack(side=RIGHT,) 
    password = tk.Entry(frame2,width=20)
    password.pack(side=RIGHT,)

    def change_art(event):
        number_var2.delete(1.0, END)
        selectedItem = tree.selection()[0]
        name = int(tree.item(selectedItem,'values')[0])
        global art_no
        art_no = int(name)-1
        file = art[art_no]
        file2=docx.Document(file)
        p=Path(file)
        number_var.set(f'标题:  {p.stem}')
        text1 = print_text(file2)
        number_var2.insert(END,text1)
    tree.bind('<Double-1>',change_art)


    l = tk.Label(frame4, bg='yellow', width=40, height=2,text='请选择模式：')
    l.pack()

    #单选窗口
    var = tk.StringVar()
    r1 = tk.Radiobutton(frame4, text='私密',variable=var, value='私密',command=print_selection)
    r1.pack()

    r2 = tk.Radiobutton(frame4, text='草稿',variable=var, value='草稿',command=print_selection)
    r2.pack()

    r3 = tk.Radiobutton(frame4, text='公开',variable=var, value='公开',command=print_selection)
    r3.pack()

    label01 = Label(frame4,text='分类')
    label01.pack()  
    category = tk.Entry(frame4,width=20)
    category.pack()
    label02 = Label(frame4,text='标签')
    label02.pack() 
    post_tag = tk.Entry(frame4,width=20)
    post_tag.pack()

    button = tk.Button(frame4,text='发送',command=post_new_article).pack()

    root.mainloop()