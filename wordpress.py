import sys
sys.path.append("E:\code\python\Lib\site-packages")
from wordpress_xmlrpc import Client, WordPressPost
from wordpress_xmlrpc.methods import posts
import docx
import glob
import pathlib 
import tkinter as tk
from tkinter import *
from tkinter import messagebox

wp = Client('http://dezhao.site/xmlrpc.php', 'admin', 'password') #地址，例如http://dezhao.site/xmlrpc.php，用户名，密码

#上传
def post_new_article(title, content,category,post_tag):
    post = WordPressPost()
    post.title = title
    post.content = content
    post.post_status = 'draft'  # 文章状态，不写默认是草稿，private表示私密的，draft表示草稿，publish表示发布
    post.terms_names = {
        'category': [category],
        'post_tag': [post_tag],
    }
    
    post.id = wp.call(posts.NewPost(post))
    return post.id
#导入文件
def search():
    f = glob.glob(r'./article/*')
    return f
#doc与docx的转换
def rename(p):
    if p.suffix == '.doc':
        p.rename(p.with_suffix('.docx'))
#输出文本内容      
def print_text(file):
    text1 = "段落数:"+str(len(file.paragraphs))
    for para in file.paragraphs:
        text1 = (text1 + '\n'+ para.text)
    return text1
#tk
def change_state():
    global cat
    cat = category.get()
    global tag
    tag = post_tag.get()
    window.destroy()

if __name__ == '__main__':
    f = search()
    for p in f:
        file=docx.Document(p)
        p = Path(p)
        
        window = tk.Tk()
        title = Label(window,text=p.stem)
        title.pack()  
        label01 = Label(window,text='分类')
        label01.pack()  
        category = tk.Entry(window,width=20)
        category.pack()
        label02 = Label(window,text='标签')
        label02.pack() 
        post_tag = tk.Entry(window,width=20)
        post_tag.pack()
        button = tk.Button(window,text='发送',command=change_state).pack()
        window.mainloop()
        
        post_id = post_new_article(p.stem,print_text(file),cat,tag) 
        print(p.stem + post_id) 
    